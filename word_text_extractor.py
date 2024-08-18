import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import os
import webbrowser
from concurrent.futures import ThreadPoolExecutor
from threading import Thread

class WordExtractorApp:
    def __init__(self, root):
        self.root = root
        self.keyword_groups = []  # 确保初始化
        self.setup_ui()

    def setup_ui(self):
        self.root.title("Word文档文本提取器")
        self.root.geometry("1000x600")  # 调整窗口大小以适应内容
        self.create_widgets()
        self.source_files = []
        self.save_folder = ""
        self.target_file = ""

    def create_widgets(self):
        self.source_label = tk.Label(self.root, text="选择要处理的Word文档：")
        self.source_label.pack(pady=10)
        self.source_button = tk.Button(self.root, text="选择文件", command=self.select_files)
        self.source_button.pack(pady=10)

        self.file_listbox_frame = tk.Frame(self.root)
        self.file_listbox_scrollbar = tk.Scrollbar(self.file_listbox_frame)
        self.file_listbox = tk.Listbox(self.file_listbox_frame, width=80, height=5, yscrollcommand=self.file_listbox_scrollbar.set)
        self.file_listbox_scrollbar.config(command=self.file_listbox.yview)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH)
        self.file_listbox_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_listbox_frame.pack(pady=10)

        self.keyword_table = tk.Frame(self.root)
        self.keyword_table.pack(pady=10)

        # 添加列标题（不包括"组别"列）
        tk.Label(self.keyword_table, text="起始关键词", anchor="center").grid(row=0, column=1, padx=5, pady=5, sticky='ew')
        tk.Label(self.keyword_table, text="结束关键词", anchor="center").grid(row=0, column=2, padx=5, pady=5, sticky='ew')
        tk.Label(self.keyword_table, text="是否包含", anchor="center").grid(row=0, column=3, padx=5, pady=5, sticky='ew')
        tk.Label(self.keyword_table, text="关键词格式", anchor="center").grid(row=0, column=4, padx=5, pady=5, sticky='ew')
        tk.Label(self.keyword_table, text="自定义命名", anchor="center").grid(row=0, column=5, padx=5, pady=5, sticky='ew')

        self.add_keyword_group()

        self.add_group_button = tk.Button(self.root, text="添加关键词组", command=self.add_keyword_group)
        self.add_group_button.pack(pady=10)

        self.format_frame = tk.Frame(self.root)
        self.format_frame.pack(pady=10)
        self.format_option = tk.IntVar(value=1)
        self.format_label = tk.Label(self.format_frame, text="选择文本保存方式：")
        self.format_label.pack(side=tk.LEFT, padx=5)
        self.radio1 = tk.Radiobutton(self.format_frame, text="按原样保留段落", variable=self.format_option, value=1)
        self.radio1.pack(side=tk.LEFT, padx=5)
        self.radio2 = tk.Radiobutton(self.format_frame, text="合并为一个段落", variable=self.format_option, value=2)
        self.radio2.pack(side=tk.LEFT, padx=5)

        self.save_option_frame = tk.Frame(self.root)
        self.save_option_frame.pack(pady=10)
        self.save_option = tk.IntVar(value=2)  # 默认选项改为新建文件保存
        self.save_option_label = tk.Label(self.save_option_frame, text="选择保存位置：")
        self.save_option_label.pack(side=tk.LEFT, padx=5)
        self.save_radio1 = tk.Radiobutton(self.save_option_frame, text="保留至已有文件", variable=self.save_option, value=1)
        self.save_radio1.pack(side=tk.LEFT, padx=5)
        self.save_radio2 = tk.Radiobutton(self.save_option_frame, text="新建文件保存", variable=self.save_option, value=2)
        self.save_radio2.pack(side=tk.LEFT, padx=5)

        self.button_frame = tk.Frame(self.root)
        self.button_frame.pack(pady=10)
        self.extract_button = tk.Button(self.button_frame, text="开始提取", command=self.start_extraction_thread)
        self.extract_button.pack(side=tk.LEFT, padx=5)
        self.open_folder_button = tk.Button(self.button_frame, text="打开保存文件夹", command=self.open_save_folder, state=tk.DISABLED)
        self.open_folder_button.pack(side=tk.LEFT, padx=5)
        self.exit_button = tk.Button(self.button_frame, text="退出", command=self.root.quit)
        self.exit_button.pack(side=tk.LEFT, padx=5)

        self.progress_label = tk.Label(self.root, text="")
        self.progress_label.pack(pady=10)

    def add_keyword_group(self):
        group_number = len(self.keyword_groups) + 1

        tk.Label(self.keyword_table, text=f"{group_number}").grid(row=group_number, column=0, padx=5, pady=5)
        keyword_start_entry = tk.Entry(self.keyword_table, width=15)
        keyword_start_entry.grid(row=group_number, column=1, padx=5, pady=5)
        keyword_end_entry = tk.Entry(self.keyword_table, width=15)
        keyword_end_entry.grid(row=group_number, column=2, padx=5, pady=5)

        contain_options = ["不含", "全包含", "仅包含起始", "仅包含结束"]
        contain_combobox = ttk.Combobox(self.keyword_table, values=contain_options, width=13)
        contain_combobox.grid(row=group_number, column=3, padx=5, pady=5)
        contain_combobox.set("仅包含起始")

        format_options = ["高亮", "无指定格式", "下划线", "加粗", "斜体"]
        format_combobox = ttk.Combobox(self.keyword_table, values=format_options, width=13)
        format_combobox.grid(row=group_number, column=4, padx=5, pady=5)
        format_combobox.set("高亮")

        custom_field_entry = tk.Entry(self.keyword_table, width=15)
        custom_field_entry.grid(row=group_number, column=5, padx=5, pady=5)

        self.keyword_groups.append({
            'keyword_start_entry': keyword_start_entry,
            'keyword_end_entry': keyword_end_entry,
            'contain_combobox': contain_combobox,
            'format_combobox': format_combobox,
            'custom_field_entry': custom_field_entry
        })

    def select_files(self):
        self.source_files = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx")])
        if self.source_files:
            self.file_listbox.delete(0, tk.END)
            for file in self.source_files:
                self.file_listbox.insert(tk.END, file)
            self.file_listbox.config(height=min(5, len(self.source_files)))
            self.source_label.config(text=f"已选择 {len(self.source_files)} 个文件")
        else:
            self.source_label.config(text="未选择文件")

    def save_options(self):
        if self.save_option.get() == 1:
            self.target_file = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")], title="选择已有文件")
            if not self.target_file:
                messagebox.showwarning("警告", "未选择已有文件，操作取消")
                return False
        elif self.save_option.get() == 2:
            self.save_folder = filedialog.askdirectory(title="选择保存文件夹", parent=self.root)
            if not self.save_folder:
                messagebox.showwarning("警告", "未选择保存文件夹，操作取消")
                return False
        return True

    def copy_paragraph_format(self, source_para, target_para):
        target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
        target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
        for run in source_para.runs:
            new_run = target_para.add_run(run.text)
            self.copy_run_format(run, new_run)

    def copy_run_format(self, source_run, target_run):
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name
        target_run.font.color.rgb = source_run.font.color.rgb
        target_run.font.highlight_color = source_run.font.highlight_color

    def start_extraction_thread(self):
        extraction_thread = Thread(target=self.extract_text)
        extraction_thread.start()

    def extract_text(self):
        keyword_groups = []
        for group in self.keyword_groups:
            keyword_start = group['keyword_start_entry'].get().strip()
            keyword_end = group['keyword_end_entry'].get().strip()
            contain_option = group['contain_combobox'].get().strip()
            keyword_format = group['format_combobox'].get().strip()
            custom_field = group['custom_field_entry'].get().strip()
            if keyword_start and keyword_end:
                keyword_groups.append({
                    'start': keyword_start,
                    'end': keyword_end,
                    'contain': contain_option,
                    'format': keyword_format,
                    'custom_field': custom_field
                })

        if not keyword_groups:
            messagebox.showwarning("警告", "请输入至少一个关键词组")
            return

        if not self.source_files:
            messagebox.showwarning("警告", "请选择至少一个文件")
            return

        if not self.save_options():
            return

        self.update_progress("提取进行中...")

        with ThreadPoolExecutor() as executor:
            results = []
            for group in keyword_groups:
                group_results = list(executor.map(self.process_file, self.source_files,
                                                  [group['start']] * len(self.source_files),
                                                  [group['end']] * len(self.source_files),
                                                  [group['contain']] * len(self.source_files),
                                                  [group['format']] * len(self.source_files),
                                                  [group['custom_field']] * len(self.source_files)))
                results.extend(group_results)

        saved_files, errors = [], []
        for result in results:
            if isinstance(result, str) and result.startswith("成功"):
                saved_files.append(result)
            else:
                errors.append(result)

        if saved_files or self.save_option.get() == 1:
            self.open_folder_button.config(state=tk.NORMAL)
            self.update_progress("提取完成！文件已保存。")
            messagebox.showinfo("完成", f"提取完成！成功 {len(saved_files)} 个文件，失败 {len(errors)} 个文件。")
        else:
            self.update_progress("没有保存任何文件。")
            messagebox.showwarning("警告", "没有保存任何文件。")

        if errors:
            error_message = "\n".join(errors)
            messagebox.showwarning("处理错误", f"以下文件处理失败：\n{error_message}")

    def process_file(self, file, keyword_start, keyword_end, contain_option, keyword_format, custom_field):
        try:
            doc = Document(file)
            extracting = False
            extracted_paragraphs = []
            for para in doc.paragraphs:
                if self.match_keyword(para, keyword_start, keyword_format):
                    extracting = True
                    if contain_option in ["全包含", "仅包含起始"]:
                        extracted_paragraphs.append(para)
                    continue
                if self.match_keyword(para, keyword_end, keyword_format) and extracting:
                    if contain_option in ["全包含", "仅包含结束"]:
                        extracted_paragraphs.append(para)
                    break
                if extracting:
                    extracted_paragraphs.append(para)

            if not extracted_paragraphs:
                return f"{file} 未提取到任何内容（未找到起始关键字或结束关键字）"

            if self.save_option.get() == 1:
                target_doc = Document(self.target_file)
                for para in extracted_paragraphs:
                    new_para = target_doc.add_paragraph()
                    self.copy_paragraph_format(para, new_para)
                    new_para._element.append(para._element)
                target_doc.save(self.target_file)
                return f"成功：{file}"
            elif self.save_option.get() == 2:
                base_name = os.path.splitext(os.path.basename(file))[0]
                save_name = f"{custom_field}-{base_name}.docx"
                target_path = os.path.join(self.save_folder, save_name)
                target_doc = Document()
                for para in extracted_paragraphs:
                    new_para = target_doc.add_paragraph()
                    self.copy_paragraph_format(para, new_para)
                    new_para._element.append(para._element)
                target_doc.save(target_path)
                return f"成功：{file}"

        except Exception as e:
            return f"{file} 处理失败: {e}"

    def match_keyword(self, para, keyword, keyword_format):
        for run in para.runs:
            if keyword in run.text:
                if keyword_format == "加粗" and not run.bold:
                    continue
                if keyword_format == "斜体" and not run.italic:
                    continue
                if keyword_format == "下划线" and not run.underline:
                    continue
                if keyword_format == "高亮" and not run.font.highlight_color:
                    continue
                return True
        return False

    def update_progress(self, message):
        self.progress_label.config(text=message)
        self.root.update_idletasks()

    def open_save_folder(self):
        if self.save_folder:
            webbrowser.open(self.save_folder)
        elif self.target_file:
            webbrowser.open(os.path.dirname(self.target_file))

if __name__ == "__main__":
    root = tk.Tk()
    app = WordExtractorApp(root)
    root.mainloop()
