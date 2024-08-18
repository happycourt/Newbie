import requests
from bs4 import BeautifulSoup
from docx import Document
import time

def fetch_news_article(url):
    # 添加用户代理
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.0.0 Safari/537.36'
    }
    response = requests.get(url, headers=headers)

    try:
        print(f"Fetching article from {url}")
        # 使用用户代理发送HTTP GET请求
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # 检查请求是否成功
        soup = BeautifulSoup(response.text, 'html.parser')

        # 以下代码不变...


        # 获取新闻标题和日期
        title_tag = soup.find('div', class_='head').find('h2')
        title = title_tag.text.strip() if title_tag else "No title"
        date_tag = soup.find('div', class_='head').find('p')
        date = date_tag.text.strip() if date_tag else "No date"

        # 获取文章简介
        intro_tag = soup.find('div', class_='content').find_all('p')[1]
        intro = intro_tag.text.strip() if intro_tag else "No intro"

        # 获取所有level的正文
        levels = ['level0', 'level1', 'level2', 'level3', 'level4', 'level5']
        article_contents = {}
        for level in levels:
            level_tag = soup.find('div', class_=level)
            if level_tag:
                content = []
                paragraphs = level_tag.find_all('p')
                for para in paragraphs:
                    if "Try the same news story at these levels:" in para.text:
                        break
                    content.append(para.text.strip())
                article_contents[level] = "\n".join(content)

        # 获取文章来源
        sources_tag = soup.find('div', class_='sources')
        sources = []
        if sources_tag:
            source_links = sources_tag.find_all('a')
            for link in source_links:
                sources.append(link['href'])

        return {
            'title': title,
            'date': date,
            'intro': intro,
            'contents': article_contents,
            'sources': sources
        }
    except requests.exceptions.RequestException as e:
        print(f"Error fetching article: {e}")
        return None

def save_to_word(news_article):
    try:
        print(f"Saving article: {news_article['title']}")
        document = Document()
        document.add_heading(news_article['title'], 0)

        document.add_paragraph(f"Date: {news_article['date']}")
        document.add_paragraph("Intro:")
        document.add_paragraph(news_article['intro'])

        for level, content in news_article['contents'].items():
            document.add_heading(level.capitalize(), level=1)
            document.add_paragraph(content)

        document.add_heading('Sources', level=1)
        for source in news_article['sources']:
            document.add_paragraph(source)

        file_name = f"{news_article['date']} {news_article['title']}.docx".replace(":", "-")
        document.save(file_name)
    except Exception as e:
        print(f"Error saving article: {e}")

def main():
    try:
        base_url = "https://breakingnewsenglish.com/"
        response = requests.get(base_url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        # 获取首页所有新闻链接
        news_links = soup.find_all('a', class_='newstory')

        for link in news_links:
            news_url = base_url + link['href']
            news_article = fetch_news_article(news_url)
            if news_article:
                save_to_word(news_article)
            time.sleep(2)  # 控制抓取频率，避免对网站造成负担
    except requests.exceptions.RequestException as e:
        print(f"Error fetching main page: {e}")

if __name__ == "__main__":
    main()
